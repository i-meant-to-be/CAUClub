import docx2pdf
import openpyxl
import docx
import time

def bDayStrStrip(input: str):
    # 생일 내 '.'를 전부 '-'으로 변경
    input = input.replace(".", "-")
    # 생일 텍스트 끝에 '-'가 있을 경우 이거 제거
    if input[-1] == "-": input = input[:-1]

    # 1. 생일이 403 (000403), 1103 (001103), 10403 (010403)
    if 3 <= len(input) <= 5:
        temp = "0" * (6 - len(input)) + input
        year, month, day = "20" + temp[0:2], temp[2:4], temp[4:]
    # 2. 생일이 990813
    elif len(input) == 6:
        year, month, day = input[0:2], input[2:4], input[4:]
        if 50 < int(year) <= 99: year = "19" + year
        else: year = "20" + year
    # 3. 생일이 99-08-13
    elif input[2] == "-" and len(input) == 8:
        year, month, day = input.split("-")
        if 50 < int(year) <= 99: year = "19" + year
        else: year = "20" + year
    # 4. 생일이 19990813
    elif len(input) == 8:
        year, month, day = input[0:4], input[4:6], input[6:]
    # 5. 생일이 1999-08-13
    elif len(input) == 10:
        year, month, day = input.split("-")
    else:
        raise Exception
    
    return f"{year}년 {int(month)}월 {int(day)}일"

def issueDoc(userName, clubName, clubMedia):
    try:
        ### 파일 불러오기
        wb_data = openpyxl.load_workbook(filename = "data.xlsx")
        wb_log = openpyxl.load_workbook(filename = "log.xlsx")
        doc = docx.Document("origin.docx")

        ### 변수 선언
        histories = []
        date = str()
        first_counter = False
        timestamp = time.localtime()

        ### 엑셀에서 기초 정보 불러오기
        for worksheet in wb_data.worksheets:
            for i in range(2, worksheet.max_row + 1):
                if worksheet["C" + str(i)].value == userName:
                    if not first_counter:
                        userId = worksheet["D" + str(i)].value
                        userDept = worksheet["E" + str(i)].value
                        bDay = bDayStrStrip(str(worksheet["H" + str(i)].value).replace(" ", ""))
                    histories.append((worksheet["A" + str(i)].value, worksheet["B" + str(i)].value, worksheet["G" + str(i)].value))

        ### 사용자가 DB에 존재하는 지 확인
        # 존재할 경우, 문서 발급 절차를 진행하고 True 반환
        if len(histories):
            # 오늘 날짜 문자열로 변환
            date = str(timestamp.tm_year) + "년 " + str(timestamp.tm_mon) + "월 " + str(timestamp.tm_mday) + "일"

            # 로그에 현재 연도 발급 대장이 없을 경우, 새로 생성
            if str(timestamp.tm_year) not in wb_log.sheetnames:
                wb_log.copy_worksheet(wb_log["template"])
                wb_log["template Copy"].title = str(timestamp.tm_year)

            ### 로그를 바탕으로 문서 번호 및 파일명 지정
            # 문서 번호를 마지막 문서 번호 + 1로 지정 
            counter = wb_log[str(timestamp.tm_year)]["A1"].value
            docNum = timestamp.tm_year * 10000 + counter + 1
            counter += 1
            wb_log[str(timestamp.tm_year)]["A1"].value = counter
            # 파일 이름 지정
            fileName = f"{docNum}_{userId}"

            ### 문서 내용 수정
            # 문서 번호
            doc.tables[0].rows[0].cells[0].text = f"제 {docNum}호"
            # 성명
            doc.tables[0].rows[2].cells[1].text = userName
            # 학번
            doc.tables[0].rows[3].cells[1].text = str(userId)
            # 학부
            doc.tables[0].rows[4].cells[1].text = userDept
            # 동아리명
            doc.tables[0].rows[5].cells[1].text = clubName
            # 동아리 매체
            doc.tables[0].rows[6].cells[1].text = clubMedia
            # 생일
            doc.tables[0].rows[7].cells[1].text = bDay
            # 날짜
            doc.tables[0].rows[16].cells[0].text = date
            # 활동 이력
            for history in histories:
                userType = history[2] if history[2] else "일반 회원"
                doc.tables[0].rows[11].cells[0].add_paragraph(f"- {history[0]}년 {history[1]}학기 / {userType}으로 활동")
            # 스타일 적용
            for i in range(2, 8): doc.tables[0].rows[i].cells[1].paragraphs[0].style = "UserData"
            for paragraph in doc.tables[0].rows[11].cells[0].paragraphs: paragraph.style = "Histories"
            doc.tables[0].rows[0].cells[0].paragraphs[0].style = "IssueNum"
            doc.tables[0].rows[16].cells[0].paragraphs[0].style = "DateString"

            # 변경된 로그 저장
            wb_log[str(timestamp.tm_year)]["B" + str(counter + 1)] = docNum
            wb_log[str(timestamp.tm_year)]["C" + str(counter + 1)] = userName
            wb_log[str(timestamp.tm_year)]["D" + str(counter + 1)] = f"{fileName}.pdf"

            # 문서 발급
            doc.save("temp.docx")
            wb_log.save("log.xlsx")
            docx2pdf.convert("temp.docx", f"./issued/{fileName}.pdf")
            return True
    except:
        print("Exception raised (on except section)")
        return False
    # 존재하지 않을 경우, False 반환
    print("Exception raised (on end of function)")
    return False

if __name__ == "__main__":
    # print("# 102: " + bDayStrStrip("102"))
    # print("# 1102: " + bDayStrStrip("1102"))
    # print("# 11102: " + bDayStrStrip("11102"))
    # print("# 011102: " + bDayStrStrip("011102"))
    # print("# 990102: " + bDayStrStrip("990102"))
    # print("# 99-01-02: " + bDayStrStrip("99-01-02"))
    # print("# 00-01-02: " + bDayStrStrip("00-01-02"))
    # print("# 00.01.02: " + bDayStrStrip("00.01.02"))
    # print("# 19990102: " + bDayStrStrip("19990102"))
    # print("# 1999-01-02: " + bDayStrStrip("1999-01-02"))
    # print("# 1999.01.02: " + bDayStrStrip("1999.01.02"))
    # print("# 1999.01.02.: " + bDayStrStrip("1999.01.02."))
    issueDoc("강시운", "비꼼", "매체")