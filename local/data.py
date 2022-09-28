from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx2pdf import convert
import openpyxl
import docx
import datetime
import re

# 각기 다른 생일 텍스트의 패턴을 통일하는 함수
# 입력: 임의의 생일 텍스트
# 출력: 패턴이 통일된 생일 텍스트
def bDayStrStrip(input: str):
    # 생일 텍스트 끝에 '-'가 있을 경우 이거 제거
    if input[-1] == "-": input = input[:-1]
    input = input.replace("00:00:00", "")
    input = input.strip()
    print("# 원본: " + str(input))
    # 1. 생일이 403 (000403), 1103 (001103), 10403 (010403)
    if 3 <= len(input) <= 5:
        temp = "0" * (6 - len(input)) + input
        year, month, day = "20" + temp[0:2], temp[2:4], temp[4:]
    # 2. 생일이 990813
    elif len(input) == 6:
        year, month, day = input[0:2], input[2:4], input[4:]
        if 50 < int(year) <= 99: year = "19" + year
        else: year = "20" + year
    # 4. 생일이 19990813
    elif len(input) == 8 and input.count("-") == 0:
        year, month, day = input[0:4], input[4:6], input[6:]
    # 5. 생일에 대시 두 개 들어감
    elif input.count("-") == 2:
        year, month, day = input.split("-")
        if int(year) < 100:
            if 50 < int(year) <= 99: year = "19" + year
            else: year = "20" + year
    year = str(year)
    month = str(month)
    day = str(day)

    if month[0] == "0": month = month[1:]
    if day[0] == "0": day = day[1:]

    return f"{str(year)}년 {str(month)}월 {str(day)}일"

# 사용자가 DB에 존재하는지 확인하는 함수
# 입력: 1) Excel 워크북, 2) 사용자 정보가 담긴 List
# 반환: 학번과 이름이 다 일치하면 0, 학번이나 이름 중 하나만 일치하면 1, 아예 없으면 2 반환
def isUserInDatabase(wb, userName, userId) -> int:
    nameCounter = False
    idCounter = False
    for worksheet in wb.worksheets:
        for i in range(2, worksheet.max_row + 1):
            if worksheet["C" + str(i)].value == userName and worksheet["D" + str(i)].value == userId: return 0
            elif worksheet["C" + str(i)].value == userName: nameCounter = True
            elif worksheet["D" + str(i)].value == userId: idCounter = True
    if nameCounter or idCounter: return 1
    else: return 2

# 사용자 데이터를 DB로부터 불러와 userData에 정리하는 함수
# 입력: 1) 사용자 이름
# 반환: 사용자 정보가 담긴 userData 리스트
# 주석: userData 정의
#   [0]: 학번
#   [1]: 이름
#   [2]: 동아리명
#   [3]: 학부
#   [4]: 생일 (Revised)
#   [5]: 매체
def getUserData(wb, userData, clubs) -> list:
    for worksheet in wb.worksheets:
            for i in range(2, worksheet.max_row + 1):
                if worksheet["C" + str(i)].value == userData[1] and worksheet["D" + str(i)].value == userData[0]:
                    userData.append(worksheet["E" + str(i)].value)
                    userData.append(bDayStrStrip(str(worksheet["H" + str(i)].value).replace(" ", "")))
                    userData.append(clubs[userData[2]])
                    return userData   

# 사용자의 활동 이력을 반환하는 함수
# 입력: 1) Excel 워크북, 2) 사용자 정보가 담긴 List
# 반환: 사용자의 활동 이력이 담긴 List
def getUserHistories(wb, userData) -> list:
    histories = list()
    for worksheet in wb.worksheets:
        firstCounter = False
        for i in range(1, worksheet.max_row + 1):
            if worksheet["C" + str(i)].value == userData[1] and str(worksheet["D" + str(i)].value) == str(userData[0]) and worksheet["F" + str(i)].value == userData[2]:
                histories.append((worksheet["A" + str(i)].value, worksheet["B" + str(i)].value, worksheet["G" + str(i)].value))
    return histories

# 문서 발급
# 입력: 1) Excel 워크북, 2) 인증서 원본 양식, 3) 활동 이력이 담긴 List, 4) 사용자 정보가 담긴 Dict
# 반환: 발급 성공 시 True, 그렇지 않을 시 False
def issueDoc(wb_data, wb_log, doc, histories, timestamp, userData) -> bool:
    try:
        # 오늘 날짜 문자열로 변환
        date = str(timestamp.tm_year) + "년 " + str(timestamp.tm_mon) + "월 " + str(timestamp.tm_mday) + "일"

        # 로그에 현재 연도 발급 대장이 없을 경우, 새로 생성
        if str(timestamp.tm_year) not in wb_log.sheetnames:
            wb_log.copy_worksheet(wb_log["template"])
            wb_log["template Copy"].title = str(timestamp.tm_year)

        ### 로그를 바탕으로 문서 번호 및 파일명 지정
        # 문서 번호를 마지막 문서 번호 + 1로 지정
        if not wb_log[str(timestamp.tm_year)]["A1"].value:
            wb_log[str(timestamp.tm_year)]["A1"].value = 0
        counter = wb_log[str(timestamp.tm_year)]["A1"].value
        docNum = timestamp.tm_year * 10000 + counter + 1
        counter += 1
        wb_log[str(timestamp.tm_year)]["A1"].value = counter
        # 파일 이름 지정
        fileName = f"{docNum}_{userData[0]}" # 학번

        ### 문서 내용 수정
        # 문서 번호
        doc.tables[0].rows[0].cells[0].text = f"제 {docNum}호"
        # 성명
        doc.tables[0].rows[2].cells[1].text = userData[0]
        # 학번
        doc.tables[0].rows[3].cells[1].text = str(userData[1])
        # 학부
        doc.tables[0].rows[4].cells[1].text = userData[3]
        # 동아리명
        doc.tables[0].rows[5].cells[1].text = userData[2]
        # 동아리 매체
        doc.tables[0].rows[6].cells[1].text = userData[5]
        # 생일
        doc.tables[0].rows[7].cells[1].text = userData[4]
        # 날짜
        doc.tables[0].rows[16].cells[0].text = date
        # 활동 이력
        for history in histories:
            userType = history[2] if history[2] else "일반 회원"
            doc.tables[0].rows[11].cells[0].add_paragraph(f"- {history[0]}년 {history[1]}학기 / {userType}으로 활동")
        # 스타일 적용
        for i in range(2, 8): doc.tables[0].rows[i].cells[1].paragraphs[0].style = "UserData"
        for paragraph in doc.tables[0].rows[11].cells[0].paragraphs: paragraph.style = "Histories"
        doc.tables[0].rows[0].cells[0].paragraphs[0].style = "IssueNum"
        doc.tables[0].rows[16].cells[0].paragraphs[0].style = "DateString"

        # 변경된 로그 저장
        wb_log[str(timestamp.tm_year)]["B" + str(counter + 1)] = docNum
        wb_log[str(timestamp.tm_year)]["C" + str(counter + 1)] = userData[0]
        wb_log[str(timestamp.tm_year)]["D" + str(counter + 1)] = f"{fileName}.pdf"
        wb_log[str(timestamp.tm_year)]["E" + str(counter + 1)] = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # 문서 발급
        doc.save("./data/temp.docx")
        wb_log.save("./data/log.xlsx")
        convert("./data/temp.docx", f"./data/issued/{fileName}.pdf")
        return f"{fileName}.pdf"
    except Exception as e:
        print(e)
        return False

if __name__ == "__main__":
    print(bDayStrStrip("1999-08-13"))
    print(bDayStrStrip("99-08-13"))
    print(bDayStrStrip("1999-8-13"))