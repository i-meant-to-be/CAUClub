from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx2pdf import convert
from docx import Document as Doc
import openpyxl as xl
import time

def issueDoc(userName, clubName, clubMedia):
    # 파일 불러오기
    wb = xl.load_workbook(filename = "Data.xlsx")
    ws = wb["Data"]
    doc = Doc("docs.docx")
    histories = list()
    date = list(map(str, [time.localtime().tm_year, time.localtime().tm_mon, time.localtime().tm_mday]))
    
    # 엑셀에서 기초 정보 불러오기 (단일 워크시트)
    # for i in range(1, ws.max_row + 1):
    #     coord = "C" + str(i)
    #     firstCounter = False
    #     if ws[coord].value == userName:
    #         if not firstCounter:
    #             userId = ws["D" + str(i)].value
    #             userDept = ws["E" + str(i)].value
    #         histories.append((ws["A" + str(i)].value, ws["B" + str(i)].value, ws["G" + str(i)].value))
    
    # 엑셀에서 기초 정보 불러오기 (여러 워크시트)
    for worksheet in wb.worksheets:
        firstCounter = False
        for i in range(1, worksheet.max_row + 1):
            if worksheet["C" + str(i)].value == userName:
                if not firstCounter:
                    userId = worksheet["D" + str(i)].value
                    userDept = worksheet["E" + str(i)].value
                histories.append((worksheet["A" + str(i)].value, worksheet["B" + str(i)].value, worksheet["G" + str(i)].value))

    # 사용자의 DB에 있으면 문서 발급
    if len(histories):
        table = doc.tables[0]
        table.rows[5].cells[0].paragraphs[0].text = f"1. 본 동연은 '{userName}' 학우가 중앙대학교 서울캠퍼스 '{clubMedia}' 동아리 '{clubName}'에서 '{len(histories)}'학기 동안 활동했음을 아래와 같이 인증합니다."

        table.rows[5].cells[0].tables[0].rows[0].cells[1].text = userName
        table.rows[5].cells[0].tables[0].rows[1].cells[1].text = userDept
        table.rows[5].cells[0].tables[0].rows[2].cells[1].text = str(userId)
        for i in range(3): table.rows[5].cells[0].tables[0].rows[i].cells[1].paragraphs[0].style = "table"
        # table.rows[5].cells[0].paragraphs[2].add_run = f"1) 이름: {userName}"
        # table.rows[5].cells[0].paragraphs[3].add_run = f"2) 학과: {userDept}"
        # table.rows[5].cells[0].paragraphs[4].add_run = f"3) 학번: {userId}"
        table.rows[8].cells[6].paragraphs[0].text = f"{date[0]}년 {date[1]}월 {date[2]}일"
        # style = doc.styles["normal"]
        # style._element.rPr.rFonts.set(qn("w:eastAsia"), '나눔스퀘어 ExtraBold')
        # style.font.color.rgb = RGBColor(0x0F, 0x7C, 0xBF)
        for history in histories:
            userType = history[2] if history[2] else "일반 회원"
            newParagraph = table.rows[5].cells[0].add_paragraph()
            run = newParagraph.add_run(f"\t    - {history[0]}년 {history[1]}학기 / {userType}으로 활동")
            run.font.color.rgb = RGBColor(0x0F, 0x7C, 0xBF)
            run.font.name = "나눔스퀘어 ExtraBold"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔스퀘어 ExtraBold")
        doc.save("new.docx")
        convert("new.docx", "new.pdf")
        return True
    # 사용자가 DB에 없으면 패스
    else:
        return False

if __name__ == "__main__":
    issueDoc("강시운", "비꼼", "자유 토론")